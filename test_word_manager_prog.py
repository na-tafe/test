# -*- coding: utf-8 -*-
"""
Created on Sat Oct 19 15:58:31 2024

@author: 61420
"""

import unittest
import os
from docx import Document
from word_manager_prog import WordManager

# Assuming WordManager is imported from the module where it is defined
# from your_module import WordManager

class TestWordManager(unittest.TestCase):

    def setUp(self):
        """Create a WordManager instance for testing."""
        self.filename = "test_document.docx"
        self.word_manager = WordManager(self.filename)
        self.image_filename = "aws.png"  # Add your image file path here
        
        

    def tearDown(self):
        """Remove the test document after each test."""
        if os.path.exists(self.filename):
            os.remove(self.filename)

    def test_create_document(self):
        """Test that a document is created successfully."""
        self.word_manager.create_document()
        self.assertTrue(os.path.exists(self.filename), "Document should be created.")

    def test_add_paragraph(self):
        """Test that a paragraph is added to the document."""
        self.word_manager.create_document()
        self.word_manager.add_paragraph()  # Add a paragraph
        
        doc = Document(self.filename)
        self.assertGreater(len(doc.paragraphs), 2, "Should have more than 2 paragraphs after adding one.")
        
        
    def test_add_picture(self):
        """Test that a picture is added to the document."""
        self.word_manager.create_document()
        self.word_manager.add_picture(self.image_filename)  # Add the test image
        
        doc = Document(self.filename)
        self.assertGreater(len(doc.inline_shapes), 0, "Document should contain at least one image.")

    def test_get_full_text(self):
        """Test that full text can be retrieved from the document."""
        self.word_manager.create_document()
        self.word_manager.add_paragraph()
        
        full_text = self.word_manager.get_full_text()
        self.assertIn("This is a paragraph under Main heading", full_text, "Full text should include the added paragraph.")

if __name__ == "__main__":
    unittest.main()
