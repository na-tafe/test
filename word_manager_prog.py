# -*- coding: utf-8 -*-  

"""
Created on Thu Oct 17 15:43:48 2024
@author: 61420
pip install python-docx  # Note to install the required library for working with Word documents
"""

from docx import Document  # Imports the Document class to create and manipulate Word documents
from docx.shared import Inches  # Imports Inches for setting image dimensions in the document
import os  # Imports os module for interacting with the operating system (e.g., file paths)

class WordManager:

    def __init__(self, filename):
        """Initialize the WordManager with the given filename."""
        self.filename = filename  # Store the filename for later use
        self.document = Document()  # Create a new Word Document instance


    def create_document(self):
        """Create a Word document with headings and paragraphs."""
        # Add headings and paragraphs to the document
        self.document.add_heading('Main Heading', level=1)  # Adds a main heading
        self.document.add_heading('Subheading', level=2)  # Adds a subheading
        self.save_document()  # Save the document after creating content

    def add_paragraph(self):
        """Add sample paragraphs to the document."""
        p_sample = (
            "Lorem ipsum dolor sit amet, consectetuer adipiscing elit. "
            "Aenean commodo ligula eget dolor. Aenean massa. Cum sociis "
            "natoque penatibus et magnis dis parturient montes, nascetur "
            "ridiculus mus. Donec quam felis, ultricies nec, pellentesque eu, "
            "pretium quis, sem. Nulla consequat massa quis enim. Donec pede justo, "
            "fringilla vel, aliquet nec, vulputate eget, arcu. In enim justo, "
            "rhoncus ut, imperdiet a, venenatis vitae, justo. Nullam dictum "
            "felis eu pede mollis pretium. Integer tincidunt."
        )
        main_heading = self.document.paragraphs[1]  # Main heading is the first paragraph
        main_heading.insert_paragraph_before("This is a paragraph under Main heading")  # Insert paragraph after the main heading
        self.document.add_paragraph(p_sample)
        

        self.save_document()  # Save the document after creating content

    def add_picture(self, image_path):
        """Add a picture to the Word document."""
        if os.path.exists(image_path):  # Check if the specified image file exists
            self.document.add_picture(image_path, width=Inches(3))  # Add the picture with a width of 3 inches
            self.save_document()  # Save the document after adding the picture
        else:
            print(f"Image not found: {image_path}")  # Notify if the image file does not exist

    def get_full_text(self):
        """Retrieve full text from the Word document."""
        full_text = []  # Initialize an empty list to hold the text
        for paragraph in self.document.paragraphs:  # Iterate through each paragraph in the document
            full_text.append(paragraph.text)  # Add the text of the paragraph to the list
        return '\n'.join(full_text)  # Join all paragraphs into a single string with line breaks

    def save_document(self):
        """Save the document to the specified filename."""
        self.document.save(self.filename)  # Save the current state of the document to the specified filename

# Example Usage
def main():
    while True:
        # Create an instance of WordManager with the filename 'example.docx'
        filename_input = input("To start the Word Manager program, enter a filename (.docx), or 'x' to exit: ")
        
        if filename_input.lower() == 'x':
            print("You are exiting the program...")
            break
        
        if filename_input.endswith('.docx'):
            word_manager = WordManager(filename_input)
            word_manager.create_document()

            q1 = input("Would you like to add paragraphs to headings? (y/n) ").lower()
            
            if q1 == 'y':
                word_manager.add_paragraph()  # Create a new document with headings and paragraphs
                print("Successfully added paragraphs.")
                
                q2 = input("Would you like to add a picture? (y/n) ").lower()
                if q2 == 'y':
                    image_path = input("Enter the image path: ")
                    word_manager.add_picture(image_path)  # Add a picture to the document
                    
                    q3 = input("Would you like to get the full text? (y/n): ").lower()
                    if q3 == 'y':
                        full_text = word_manager.get_full_text()
                        print("Full text of the document:\n")
                        print(full_text)  # Output the full text of the document
                    elif q3 == 'n':
                        print("You are exiting the program...")
                        break
                    else:
                        print("Invalid input. Please enter 'y' or 'n'.")
                elif q2 == 'n':
                    print("You are exiting the program...")
                    break
                else:
                    print("Invalid input. Please enter 'y' or 'n'.")
            elif q1 == 'n':
                print("No paragraphs added. You are exiting the program.")
                break  # Exit the loop if the user chooses not to create a document
            else:
                print("Invalid input. Please enter 'y' or 'n'.")
        else:
            print("Invalid filename. Please ensure it ends with .docx.")

if __name__ == "__main__":  # Entry point of the program
    main()
